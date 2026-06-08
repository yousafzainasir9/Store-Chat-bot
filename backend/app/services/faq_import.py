"""Parse an uploaded FAQ document into content items.

Supports out of the box (no extra deps): Markdown/text (split by ``##``
sections, like the seed KB) and CSV (question/answer rows). PDF and DOCX are
parsed via lazily-imported optional libraries (``pip install ".[docs]"``); if the
library is missing a clear error is raised.

Returns a list of :class:`ParsedFaq` that the admin endpoint feeds to
:class:`ContentService`, which indexes each for retrieval.
"""

from __future__ import annotations

import csv
import io
from dataclasses import dataclass
from pathlib import PurePath

from app.services.seed import parse_markdown_sections

_MAX_ITEMS = 1000


class UnsupportedDocument(Exception):
    """Raised for an unsupported file type."""


class DocumentParseError(Exception):
    """Raised when a document cannot be parsed (e.g. missing optional library)."""


@dataclass(slots=True)
class ParsedFaq:
    """One FAQ entry extracted from an uploaded document."""

    title: str
    body: str
    category: str = "FAQ"


def _decode(data: bytes) -> str:
    for enc in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _from_markdown(text: str, stem: str) -> list[ParsedFaq]:
    sections = parse_markdown_sections(text)
    if sections:
        return [ParsedFaq(title=title, body=content) for title, content in sections]
    cleaned = text.strip()
    return [ParsedFaq(title=stem, body=cleaned)] if cleaned else []


def _from_csv(text: str) -> list[ParsedFaq]:
    rows = list(csv.reader(io.StringIO(text)))
    if not rows:
        return []
    # Detect a header row naming the columns; otherwise assume col0=Q, col1=A.
    header = [c.strip().lower() for c in rows[0]]
    q_idx, a_idx, start = 0, 1, 0
    if any(h in {"question", "title", "q"} for h in header) and any(
        h in {"answer", "body", "a"} for h in header
    ):
        q_idx = next(i for i, h in enumerate(header) if h in {"question", "title", "q"})
        a_idx = next(i for i, h in enumerate(header) if h in {"answer", "body", "a"})
        start = 1
    out: list[ParsedFaq] = []
    for row in rows[start:]:
        if len(row) <= max(q_idx, a_idx):
            continue
        title, body = row[q_idx].strip(), row[a_idx].strip()
        if title and body:
            out.append(ParsedFaq(title=title, body=body))
    return out


def _from_pdf(data: bytes, stem: str) -> list[ParsedFaq]:
    try:
        from pypdf import PdfReader  # lazy import (optional `docs` extra)
    except ImportError as exc:  # pragma: no cover - depends on optional install
        raise DocumentParseError(
            "PDF import requires the 'docs' extra: pip install '.[docs]'"
        ) from exc
    reader = PdfReader(io.BytesIO(data))
    text = "\n".join((page.extract_text() or "") for page in reader.pages).strip()
    return [ParsedFaq(title=stem, body=text)] if text else []


def _from_docx(data: bytes, stem: str) -> list[ParsedFaq]:
    try:
        import docx  # lazy import (optional `docs` extra)
    except ImportError as exc:  # pragma: no cover - depends on optional install
        raise DocumentParseError(
            "DOCX import requires the 'docs' extra: pip install '.[docs]'"
        ) from exc
    document = docx.Document(io.BytesIO(data))
    text = "\n".join(p.text for p in document.paragraphs).strip()
    return [ParsedFaq(title=stem, body=text)] if text else []


def parse_faq_document(filename: str, data: bytes) -> list[ParsedFaq]:
    """Parse ``data`` (an uploaded file) into FAQ items based on its extension."""
    suffix = PurePath(filename).suffix.lower()
    stem = PurePath(filename).stem or "FAQ"
    if suffix in {".md", ".markdown", ".txt", ""}:
        items = _from_markdown(_decode(data), stem)
    elif suffix == ".csv":
        items = _from_csv(_decode(data))
    elif suffix == ".pdf":
        items = _from_pdf(data, stem)
    elif suffix == ".docx":
        items = _from_docx(data, stem)
    else:
        raise UnsupportedDocument(f"Unsupported file type: {suffix or 'unknown'}")
    return items[:_MAX_ITEMS]
